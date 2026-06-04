from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "RentFlow_Demo_Readiness_Report_2026-06-04.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths=None, header_fill="E8EEF5"):
    table.style = "Table Grid"
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for idx, cell in enumerate(row.cells):
            if widths:
                cell.width = Inches(widths[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for m, v in [("top", "80"), ("bottom", "80"), ("start", "120"), ("end", "120")]:
                node = margins.find(qn(f"w:{m}"))
                if node is None:
                    node = OxmlElement(f"w:{m}")
                    margins.append(node)
                node.set(qn("w:w"), v)
                node.set(qn("w:type"), "dxa")
            if row_idx == 0:
                set_cell_shading(cell, header_fill)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + level * 0.2)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)
    return p


def add_para(doc, text="", bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


def add_score_table(doc):
    rows = [
        ("Auth/session/role", "20", "16", "Route guard/proxy tests pass, auth/session logs show login/session behavior; thiếu full browser invalid/logout/back-button matrix."),
        ("Customer booking", "20", "15", "Backend booking/idempotency/concurrency evidence mạnh; frontend booking unit tests có. Chưa có full E2E customer create/cancel/reload evidence."),
        ("Host flow", "15", "11", "Host dashboard/listing/booking coverage có ở tests/docs; thiếu browser E2E host approve/reject/ownership smoke."),
        ("Admin flow", "15", "11", "Admin pages/build smoke có evidence; backend/admin coverage gián tiếp. Chưa đủ action-level browser regression."),
        ("Cancellation/payment", "15", "13", "Cancellation/payment backend gate rất mạnh, CoreBank/manual transfer contract đã green; thiếu demo-path E2E qua UI + provider/Mailpit."),
        ("Email/Mailpit", "5", "2", "Có email service tests/docs, nhưng chưa có Mailpit regression matrix trong evidence vừa thấy."),
        ("Responsive/UI", "5", "2", "Có frontend tests và smoke route, nhưng chưa thấy mobile/tablet screenshot/regression matrix."),
        ("Regression/build/tests", "5", "5", "Backend unit/integration, frontend Vitest và production build đều có bằng chứng pass gần đây."),
    ]
    table = doc.add_table(rows=1, cols=5)
    style_table(table, widths=[1.55, 0.55, 0.55, 0.75, 3.1])
    headers = ["Hạng mục", "Max", "Điểm", "Mức", "Lý do"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    total = 0
    for area, max_score, score, note in rows:
        total += int(score)
        row = table.add_row().cells
        rating = "Tốt" if int(score) / int(max_score) >= 0.8 else "Rủi ro"
        vals = [area, max_score, score, rating, note]
        for idx, val in enumerate(vals):
            set_cell_text(row[idx], val, bold=(idx == 0))
    return total


def add_fix_table(doc):
    data = [
        ("1", "P1", "E2E role/session", "Chạy lại Browser regression cho register/login/logout/reload/back-button, customer/host/admin route guard, API missing/invalid token.", "Nếu fail redirect/session thì fix ngay trước demo."),
        ("2", "P1", "Customer booking happy path", "Tạo booking mới qua UI, xem list/detail, reload, duplicate submit, invalid dates, cancel allowed state.", "Đây là flow portfolio chính, phải có video/screenshot pass."),
        ("3", "P1", "Host approve/reject", "Host xem booking request, approve/reject booking test, kiểm tra customer detail và status sau reload.", "Không demo nếu host workspace không pass."),
        ("4", "P1", "Payment/cancellation UI", "CoreBank hoặc bank-transfer demo từ booking payment tới authorize/void/refund/cancel state; kiểm tra idempotency/double cancel.", "Backend green nhưng cần chứng minh UI tích hợp thật."),
        ("5", "P2", "Mailpit + responsive", "Kiểm tra email registration/booking/cancellation/host decision nếu flow có; chụp 390px cho home/listing/booking/admin/host.", "Không nhất thiết P0, nhưng portfolio nhìn sẽ kém nếu vỡ layout/email chết."),
    ]
    table = doc.add_table(rows=1, cols=5)
    style_table(table, widths=[0.35, 0.55, 1.45, 2.75, 2.25])
    headers = ["#", "Ưu tiên", "Khu vực", "Việc cần làm", "Tiêu chí quyết định"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for item in data:
        row = table.add_row().cells
        for idx, val in enumerate(item):
            set_cell_text(row[idx], val, bold=(idx in [0, 1, 2]))


def add_evidence_table(doc):
    data = [
        ("docs/release-gate-evidence.md", "2026-06-03", "Frontend proxy targeted: 17 tests pass; full frontend Vitest: 160 tests pass; production build pass."),
        ("docs/release-gate-evidence.md", "2026-06-02", "Backend unit: 440 pass; integration final: 183 pass; targeted booking/availability: 26 pass; frontend test/build pass."),
        ("docs/release-gate-evidence.md", "2026-05-30", "Admin/payment browser smoke against mocked API; frontend build/test pass. Limit: backend infra không chạy trong session này."),
        ("frontend.log", "2026-06-04", "Local Next.js started, public listings/login/admin routes hit; admin session endpoint trả 200 sau login. Đây chỉ là log smoke, không thay thế full regression report."),
        ("Attachment pasted-text.txt", "2026-06-04", "Nội dung là checklist/yêu cầu QA, không phải actual result matrix. Vì vậy các flow thiếu evidence bị trừ điểm."),
    ]
    table = doc.add_table(rows=1, cols=3)
    style_table(table, widths=[2.0, 1.0, 3.5])
    headers = ["Nguồn", "Ngày", "Bằng chứng / giới hạn"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for item in data:
        row = table.add_row().cells
        for idx, val in enumerate(item):
            set_cell_text(row[idx], val, bold=(idx == 0))


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("RentFlow Demo/Release Readiness Report")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(10)
    r = subtitle.add_run("Ngày đánh giá: 2026-06-04 | Repo: Car-Rentail | Commit: 216c766 | Branch: main")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("555555")

    add_heading(doc, "1. Kết luận nhanh", 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Điểm readiness: 75/100 - CONDITIONAL PASS.")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("7A5A00")
    p.add_run(" Không nên đưa vào portfolio demo công khai ngay nếu chưa rerun full E2E. Có thể demo nội bộ hoặc demo có kịch bản cố định, nhưng trước khi quay/đăng portfolio nên xác nhận lại customer booking, host approve/reject, payment/cancellation UI, Mailpit và responsive.")

    add_para(doc, "Nói thẳng: hiện chưa có bug P0/P1 được chứng minh trong evidence đang có, nhưng cũng chưa có đủ evidence để khẳng định full app đã sẵn sàng portfolio-grade. Rủi ro lớn nhất không phải build/test, mà là thiếu full browser regression tích hợp backend thật.")

    add_heading(doc, "2. Bảng điểm theo hạng mục", 1)
    total = add_score_table(doc)
    add_para(doc, f"Tổng điểm: {total}/100. Cách chấm ưu tiên bằng chứng thực tế: automated gate mạnh được cộng cao; flow thiếu E2E/Mailpit/responsive evidence bị trừ dù code có thể đã đúng.")

    add_heading(doc, "3. Verdict demo", 1)
    add_bullet(doc, "Portfolio demo công khai: chưa nên đưa ngay. Cần rerun và ghi evidence cho 5 nhóm ở mục fix order.")
    add_bullet(doc, "Curated local demo: có thể demo nếu giới hạn script vào các flow đã smoke/test và chuẩn bị dữ liệu trước.")
    add_bullet(doc, "Release local/internal: conditional pass, vì build và automated regression đang rất tốt.")

    add_heading(doc, "4. Evidence đã dùng", 1)
    add_evidence_table(doc)

    add_heading(doc, "5. Điểm mạnh đã đủ tự tin", 1)
    add_bullet(doc, "Regression/build/tests rất tốt: backend unit/integration và frontend test/build đều có bằng chứng pass gần đây.")
    add_bullet(doc, "Auth route guard/proxy đã được test sau migration Next.js proxy.ts, giảm rủi ro routing/session ở frontend.")
    add_bullet(doc, "Booking/cancellation/payment backend có coverage mạnh, đặc biệt transaction hardening, idempotency, cancellation, provider mutation safety.")
    add_bullet(doc, "Admin/payment UI ít nhất đã có test và smoke route, dù một phần evidence cũ dùng mocked API.")

    add_heading(doc, "6. Rủi ro còn lại trước portfolio", 1)
    add_bullet(doc, "Thiếu full browser report thật theo checklist đính kèm: guest/auth/customer/host/admin/payment/email/security/responsive.")
    add_bullet(doc, "Mailpit chưa có matrix trigger-recipient-subject-link works, nên email chỉ đạt partial.")
    add_bullet(doc, "Responsive chưa có evidence ở mobile 390px/tablet cho dashboard, table, form, modal.")
    add_bullet(doc, "Customer booking và host approve/reject là flow kể chuyện chính của portfolio, nhưng evidence hiện tại chưa đủ E2E tích hợp backend thật.")
    add_bullet(doc, "Payment/cancellation backend mạnh, nhưng portfolio cần thấy UI chuyển trạng thái đúng sau reload và sau duplicate/double action.")

    add_heading(doc, "7. Recommended fix/rerun order", 1)
    add_fix_table(doc)

    add_heading(doc, "8. Top 5 tests nên rerun sau fix", 1)
    add_bullet(doc, "Auth/session/role Browser regression: register/login/logout/reload/back-button, role redirect, protected API missing/invalid token.")
    add_bullet(doc, "Customer booking E2E: listing detail -> valid booking -> booking detail/list -> invalid dates -> duplicate submit -> cancel.")
    add_bullet(doc, "Host E2E: host dashboard -> booking requests -> approve/reject -> ownership forbidden checks.")
    add_bullet(doc, "Payment/cancellation E2E: CoreBank/manual transfer authorize, cancel, void/refund/penalty state, double cancel/idempotency.")
    add_bullet(doc, "Mailpit + responsive: email triggers/link behavior, mobile 390px screenshots for public/customer/host/admin.")

    add_heading(doc, "9. Final answer", 1)
    add_para(doc, "Chưa nên đưa vào portfolio demo công khai ngay. Hãy fix hoặc ít nhất rerun chứng minh các flow P1 ở trên trước. Nếu tất cả 5 nhóm rerun pass, điểm có thể tăng từ 75 lên khoảng 88-92 và lúc đó đủ đẹp để đưa vào portfolio demo.")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = footer.add_run("RentFlow readiness report - generated 2026-06-04")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string("777777")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
