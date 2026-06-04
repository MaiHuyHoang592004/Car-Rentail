from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "RentFlow_Postfix_Full_Regression_Readiness_Report_2026-06-04.docx"


SCORES = [
    ("Auth/session/role", 20, 17, "PASS có lưu ý", "JWT, login/logout, 401/403 và role guard tốt; email verification local đang chặn flow tự nhiên."),
    ("Customer booking", 20, 16, "PARTIAL PASS", "API create/list/detail/idempotency/cancel pass sau khi setup verified; UI booking create chưa chạy trọn vì mail verify bị block."),
    ("Host flow", 15, 9, "PARTIAL", "Host auth/API listing/booking pass; approve/reject chưa test được vì booking chưa vào PENDING_HOST_APPROVAL."),
    ("Admin flow", 15, 8, "FAIL UI", "Admin login/dashboard pass, backend admin API 200; UI /admin/users và /admin/listings báo lỗi load dữ liệu."),
    ("Cancellation/payment", 15, 11, "PARTIAL PASS", "Cancel preview/cancel/idempotent cancel pass; bank transfer QR authorize pass; capture/void/refund/CoreBank chưa được xác nhận."),
    ("Email/Mailpit", 5, 1, "FAIL", "Mailpit reachable nhưng inbox 0; app local đang tắt mail nên không verify được user qua email."),
    ("Responsive/UI", 5, 2, "FAIL nhỏ", "Login mobile ổn; home/listings overflow ngang ở mobile 390px; admin data UI fail."),
    ("Regression/build/tests", 5, 5, "PASS evidence", "Release-gate evidence gần nhất có backend/frontend tests/build pass; chưa rerun full suite trong phiên Browser này."),
]


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Aptos Display"
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_note(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)
    if widths:
        for row in t.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return t


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RentFlow Demo Readiness Report")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(17, 24, 39)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Post-fix full regression assessment | 2026-06-04").italic = True

    total = sum(score for _, _, score, _, _ in SCORES)
    max_total = sum(max_score for _, max_score, _, _, _ in SCORES)
    verdict = doc.add_paragraph()
    verdict.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = verdict.add_run(f"Final score: {total}/{max_total} - FAIL for portfolio demo readiness")
    vr.bold = True
    vr.font.size = Pt(14)
    vr.font.color.rgb = RGBColor(185, 28, 28)

    add_note(
        doc,
        "Hiện tại chưa nên đưa vào portfolio demo công khai. Backend/API core khá ổn, nhưng demo UI còn lỗi admin load dữ liệu, email verification/Mailpit không chạy ở local, host approval/payment state chưa test được trọn vòng, và có overflow mobile dễ bị nhìn thấy.",
        "Kết luận thẳng: ",
    )

    add_heading(doc, "1. Test Environment")
    table(
        doc,
        ["Hạng mục", "Giá trị / kết quả"],
        [
            ["Frontend", "http://localhost:3000 - reachable, public routes render được."],
            ["Backend", "http://localhost:8087 - /actuator/health = UP."],
            ["Mailpit", "http://localhost:8025 - API reachable, message count = 0 sau test."],
            ["Database/infra", "Docker containers rentflow-postgres, rentflow-redis, rentflow-mailpit đang chạy."],
            ["Browser", "In-app Browser dùng để smoke public/customer/admin/host và responsive mobile 390px."],
            ["Test data", "Chỉ tạo tài khoản/dữ liệu prefix qa-postfix-*; đã setup email_verified=true cho user qa-postfix để tiếp tục booking vì Mailpit không phát mail."],
        ],
        [4.0, 12.0],
    )

    add_heading(doc, "2. Readiness Score")
    table(
        doc,
        ["Nhóm", "Điểm", "Trạng thái", "Lý do chấm"],
        [[name, f"{score}/{max_score}", status, reason] for name, max_score, score, status, reason in SCORES],
        [4.0, 2.0, 3.0, 7.2],
    )

    add_heading(doc, "3. Fix Verification Matrix")
    table(
        doc,
        ["ID", "Kỳ vọng", "Kết quả", "Evidence"],
        [
            ["AUTH-01", "Guest public route truy cập được, private route yêu cầu login.", "PASS", "/, /listings render; /me/bookings cần session; missing token /api/v1/users/me = 401."],
            ["AUTH-02", "Role guard đúng giữa customer/host/admin.", "PASS", "Customer vào /host/dashboard hoặc /admin bị /forbidden; customer gọi /api/v1/admin/users = 403; admin API = 200."],
            ["BOOK-01", "Customer có thể tạo booking hợp lệ.", "PARTIAL PASS", "API tạo booking HELD pass sau khi user được set verified; flow tự nhiên bị block bởi EMAIL_NOT_VERIFIED vì Mailpit không có mail."],
            ["BOOK-02", "Booking idempotency ổn.", "PASS", "Gửi lại cùng idempotency key trả lại cùng booking id e5e18b92-da1f-40a4-b0ba-bf252ff6ba2b."],
            ["ADMIN-01", "Admin users/listings UI load dữ liệu.", "FAIL", "UI /admin/users và /admin/listings hiện alert lỗi tải dữ liệu; backend API trực tiếp trả 200."],
            ["PAY-01", "Payment bank transfer/QR không tự mark paid.", "PASS", "Authorize với VCB tạo payment PENDING_TRANSFER; booking vẫn HELD, đúng payment rules."],
            ["MAIL-01", "Verification/notification email xuất hiện trong Mailpit.", "FAIL", "Mailpit API total = 0; application local default rentflow.mail.enabled=false."],
            ["RESP-01", "Không overflow ngang mobile.", "FAIL", "Viewport mobile report clientWidth=375, scrollWidth=383 ở / và /listings."],
        ],
        [2.2, 5.0, 3.0, 6.0],
    )

    add_heading(doc, "4. Feature Coverage Matrix")
    table(
        doc,
        ["Area", "Đã xác nhận", "Chưa xác nhận / rủi ro"],
        [
            ["Auth/session/role", "Login admin/customer/host; logout; API 401/403; route guard customer/admin/host.", "Refresh/back cache và expired refresh-token không rerun sâu trong phiên này."],
            ["Customer booking", "Register/login/list/detail/create valid booking, duplicate idempotency, cancel preview, cancel.", "UI booking create end-to-end chưa hoàn tất vì user verify mail bị block."],
            ["Host flow", "Host login, host listing API, host booking API; host route shell render.", "Approve/reject chưa test được do không có booking PENDING_HOST_APPROVAL."],
            ["Admin flow", "Admin login, dashboard shell, backend admin APIs 200.", "Admin users/listings UI fail data load, ảnh hưởng trực tiếp demo."],
            ["Payment/cancellation", "Manual bank transfer instruction generated; no false paid status; cancellation idempotent.", "CoreBank, capture, void, refund, reconciliation chưa chạy end-to-end."],
            ["Email/Mailpit", "Mailpit service reachable.", "Không có email nào được gửi; verification không demo được."],
            ["Responsive/UI", "Login mobile OK.", "Home/listings overflow ngang; admin/host mobile authenticated chưa cover đủ."],
        ],
        [3.6, 6.4, 6.2],
    )

    add_heading(doc, "5. Critical Findings")
    table(
        doc,
        ["ID", "Severity", "Finding", "Impact", "Suggested fix"],
        [
            ["RF-QA-001", "P1", "Admin UI data load fail ở /admin/users và /admin/listings dù backend admin API trả 200.", "Demo admin bị hỏng ngay khi người xem mở danh sách users/listings.", "Kiểm tra frontend API client/proxy/session cho admin pages; thêm browser test cho admin tables."],
            ["RF-QA-002", "P1", "Email verification/Mailpit không hoạt động local; user mới bị EMAIL_NOT_VERIFIED khi booking.", "Customer demo tự nhiên bị chặn, phải can thiệp DB để tiếp tục flow.", "Bật Mailpit cho local/demo hoặc seed verified demo user rõ ràng; verify resend/click link."],
            ["RF-QA-003", "P2", "Không test được host approve/reject và payment state sau transfer vì sandbox confirmation/CoreBank chưa khả dụng.", "Không chứng minh được vòng booking -> payment -> host decision -> final state.", "Bật sandbox transfer confirmation hoặc chạy CoreBank demo path; seed booking đúng state cho host flow."],
            ["RF-QA-004", "P3", "Overflow ngang mobile ở home/listings.", "Portfolio demo mobile nhìn thiếu polish.", "Fix layout/card/filter width; thêm Playwright/mobile viewport check."],
        ],
        [2.2, 1.8, 4.8, 4.0, 4.6],
    )

    add_heading(doc, "6. State Transition Matrix")
    table(
        doc,
        ["Transition", "Observed", "Result"],
        [
            ["Register -> unverified customer", "User tạo thành công; emailVerified=false.", "PASS"],
            ["Unverified customer -> create booking", "API trả 403 EMAIL_NOT_VERIFIED.", "PASS rule, FAIL demo setup"],
            ["Verified customer -> create booking", "Booking HELD, totalPrice 1,400,000 VND.", "PASS"],
            ["HELD booking -> authorize bank transfer", "Payment PENDING_TRANSFER; booking vẫn HELD.", "PASS"],
            ["HELD booking -> cancel", "Booking CANCELLED; duplicate cancel idempotent.", "PASS"],
            ["PENDING_TRANSFER -> PENDING_HOST_APPROVAL", "Không thực hiện được vì transfer confirmation disabled.", "NOT TESTED"],
            ["PENDING_HOST_APPROVAL -> approved/rejected", "Không có booking ở state phù hợp.", "NOT TESTED"],
        ],
        [5.2, 7.0, 3.4],
    )

    add_heading(doc, "7. API Regression Matrix")
    table(
        doc,
        ["Endpoint / action", "Result", "Notes"],
        [
            ["POST /api/v1/auth/login admin/customer/host", "PASS", "Tokens nhận được, roles đúng."],
            ["GET /api/v1/users/me", "PASS", "200 với token hợp lệ; 401 khi thiếu token."],
            ["GET /api/v1/admin/users", "PASS API", "Admin token trả 200; customer token trả 403."],
            ["GET /api/v1/admin/listings", "PASS API", "Admin token trả 200."],
            ["GET /api/v1/listings", "PASS", "Public endpoint trả danh sách listing."],
            ["POST /api/v1/bookings", "PASS/PARTIAL", "Pass khi customer verified; unverified bị 403 đúng rule."],
            ["GET /api/v1/bookings/me và detail", "PASS", "Customer thấy booking của mình."],
            ["POST cancel preview/cancel", "PASS", "Preview eligible true; cancel idempotent."],
            ["GET /api/v1/payment-banks", "PASS auth rule", "401 nếu không auth; 200 với customer token."],
            ["POST authorize bank transfer", "PASS", "Tạo transfer instruction, payment PENDING_TRANSFER."],
            ["GET /api/v1/host/listings, /host/bookings", "PASS API", "Host token trả dữ liệu."],
        ],
        [6.0, 2.6, 7.2],
    )

    add_heading(doc, "8. Mailpit Matrix")
    table(
        doc,
        ["Check", "Result", "Comment"],
        [
            ["Mailpit service reachable", "PASS", "HTTP API trả 200."],
            ["Registration email", "FAIL", "Message count vẫn 0."],
            ["Verification link usable", "BLOCKED", "Không có raw token/email để click; DB chỉ lưu token_hash."],
            ["Booking/cancel notification email", "FAIL/BLOCKED", "Không thấy message sau booking/cancel."],
        ],
        [5.0, 3.0, 8.0],
    )

    add_heading(doc, "9. Regression Command Result")
    table(
        doc,
        ["Command / evidence", "Result", "Source"],
        [
            ["mvn test", "PASS 440 tests", "docs/release-gate-evidence.md, 2026-06-02"],
            ["mvn verify -Pintegration-tests", "PASS 183 integration tests", "docs/release-gate-evidence.md, 2026-06-02"],
            ["frontend pnpm test", "PASS 31 files / 160 tests", "docs/release-gate-evidence.md, 2026-06-03"],
            ["frontend pnpm build", "PASS", "docs/release-gate-evidence.md, 2026-06-03"],
            ["Current session full rerun", "NOT RERUN", "Phiên này ưu tiên Browser/API regression theo checklist; không claim full suite mới."],
        ],
        [5.8, 3.0, 7.0],
    )

    add_heading(doc, "10. Recommended Fix Order")
    for item in [
        "Fix P1 admin UI data load for /admin/users and /admin/listings; verify backend 200 path reaches table/empty state in browser.",
        "Enable Mailpit/local email verification for demo, or provide seeded verified customer accounts and mark email demo limitation explicitly.",
        "Provide a complete host decision path: enable sandbox transfer confirmation or CoreBank demo, then test HELD -> PENDING_HOST_APPROVAL -> approve/reject -> final state.",
        "Fix mobile horizontal overflow on home/listings and add viewport regression check.",
        "Rerun mvn test, mvn verify -Pintegration-tests, frontend pnpm test, frontend pnpm build, then repeat Browser smoke.",
    ]:
        doc.add_paragraph(item, style="List Number")

    add_heading(doc, "11. Final Verdict")
    add_note(
        doc,
        "Không nên đưa vào portfolio demo công khai ở trạng thái hiện tại. Điểm 69/100 phản ánh rằng phần API/backend đủ mạnh để demo có kiểm soát, nhưng portfolio demo cần tự chạy mượt qua UI. Trước khi publish/record demo, tối thiểu phải fix admin UI load, Mailpit/email verification, host approval/payment progression, và mobile overflow.",
        "Verdict: ",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"Generated {datetime.now().strftime('%Y-%m-%d %H:%M')} | RentFlow QA readiness"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
